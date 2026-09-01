"""
Template para la Generación de Manuales de MMCall Academy
Este script sirve como base para estructurar nuevos manuales en formato Word y PDF
respetando el diseño estándar y la identidad visual de la empresa.

Estilo verificado directamente contra dos documentos ya publicados y de familias de
producto distintas (SENSORES/Sensor portillon/S_AlertaPermitetral.docx y
T02/MANUAL T02 PAGER.docx, en REPOSITORIO MMCALL ACADEMY/SyD/) - no es una
interpretación libre. Los dos coinciden byte a byte en fuente, colores y estructura,
así que es el estándar real de la empresa, no una particularidad de un documento.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys

# Paleta real (verificada, no la de una interpretación de la guía de marca)
C_RED = RGBColor(0xE3, 0x06, 0x13)      # Título principal, encabezado de tabla
C_LILA = RGBColor(0x9C, 0x89, 0xF5)     # "ACADEMY", eyebrow superior
C_GRAY_DARK = RGBColor(0x40, 0x40, 0x40)   # "MMCALL", encabezados de sección
C_GRAY_MED = RGBColor(0x7F, 0x8C, 0x8D)    # Pies de página / cabecera pages 2+ (igual en ambos)
C_GRAY_LIGHT = RGBColor(0x80, 0x80, 0x80)  # Subtítulo de portada, sello
C_BODY = RGBColor(0x33, 0x33, 0x33)     # Texto de cuerpo
C_TABLE_ALT = "F2F2F2"                  # Fila alterna de tabla (hex, para shading)
C_TABLE_WHITE = "FFFFFF"

FONT_BODY = "Calibri"
FONT_EYEBROW = "Consolas"


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def _brand_run(paragraph, text, color, size=None, bold=True):
    r = paragraph.add_run(text)
    r.font.name = FONT_BODY
    r.font.bold = bold
    r.font.color.rgb = color
    if size:
        r.font.size = size
    return r

def create_document_base(output_docx_name, device_name, subtitle_text):
    doc = docx.Document()

    # 1. Configurar márgenes estándar (1.0 pulgada) - verificado igual en los reales
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.different_first_page_header_footer = True

        # Cabecera (páginas 2+): "MMCALL" gris + "ACADEMY" lila + " | [Título]" gris claro
        # - las tres partes en tonos distintos, no un solo string de un color.
        header = section.header
        header_table = header.add_table(1, 2, Inches(6.5))
        header_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        cell_l = header_table.cell(0, 0)
        cell_r = header_table.cell(0, 1)

        p_head = cell_r.paragraphs[0]
        p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _brand_run(p_head, "MMCALL ", C_GRAY_DARK, Pt(10))
        _brand_run(p_head, "ACADEMY", C_LILA, Pt(10))
        r_doc = p_head.add_run(f"  |  {device_name}")
        r_doc.font.name = FONT_BODY
        r_doc.font.size = Pt(9.5)
        r_doc.font.color.rgb = C_GRAY_MED

        # Pie de página (páginas 2+) - igual en ambos documentos reales
        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_foot = p_foot.add_run("S&D Mmcall Paging Services LTDA.")
        r_foot.font.name = FONT_BODY
        r_foot.font.size = Pt(9)
        r_foot.font.color.rgb = C_GRAY_MED

    # Fuente normal: Calibri, no Arial
    style_normal = doc.styles['Normal']
    style_normal.font.name = FONT_BODY
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = C_BODY

    # --- PORTADA ---
    p_aca = doc.add_paragraph()
    p_aca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_aca = p_aca.add_run("PROGRAMA DE CAPACITACIÓN ACADÉMICA")
    r_aca.font.name = FONT_EYEBROW  # monoespaciada, no Arial - así es en los dos reales
    r_aca.font.size = Pt(9)
    r_aca.font.bold = True
    r_aca.font.color.rgb = C_LILA
    p_aca.paragraph_format.space_before = Pt(36)
    p_aca.paragraph_format.space_after = Pt(6)

    p_mmc = doc.add_paragraph()
    p_mmc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _brand_run(p_mmc, "MMCALL ", C_GRAY_DARK, Pt(24))
    _brand_run(p_mmc, "ACADEMY", C_LILA, Pt(24))
    p_mmc.paragraph_format.space_after = Pt(24)

    # Logotipo real (imagenes/logo_mmcall_nobg.png relativo al CWD del script que llama)
    logo_path = "imagenes/logo_mmcall_nobg.png"
    if os.path.exists(logo_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(logo_path, width=Inches(2.5))
        p_img.paragraph_format.space_after = Pt(24)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(device_name)
    r_title.font.name = FONT_BODY
    r_title.font.size = Pt(26)
    r_title.font.bold = True
    r_title.font.color.rgb = C_RED  # rojo, no azul marino
    p_title.paragraph_format.space_after = Pt(6)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(subtitle_text)
    r_sub.font.name = FONT_BODY
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = C_GRAY_LIGHT
    p_sub.paragraph_format.space_after = Pt(48)

    p_sello = doc.add_paragraph()
    p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sello = p_sello.add_run("Sello de Calidad MMCall Academy")
    r_sello.font.name = FONT_BODY
    r_sello.font.size = Pt(10)
    r_sello.font.italic = True
    r_sello.font.bold = False          # sin negrita
    r_sello.font.color.rgb = C_GRAY_LIGHT  # gris, no verde - así es en los dos reales
    p_sello.paragraph_format.space_after = Pt(24)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Soporte y Desarrollo (S&D) - MMCall Paging Services LTDA. - Chile")
    r_meta.font.name = FONT_BODY
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = C_GRAY_MED

    doc.add_page_break()

    # Repetición del título en rojo al iniciar la página 2 - presente en ambos
    # documentos reales, antes de la primera sección de contenido.
    p_rep = doc.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rep = p_rep.add_run(device_name)
    r_rep.font.name = FONT_BODY
    r_rep.font.size = Pt(18)
    r_rep.font.bold = True
    r_rep.font.color.rgb = C_RED
    p_rep.paragraph_format.space_after = Pt(18)

    return doc

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT_BODY
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_GRAY_DARK  # gris oscuro, no azul marino
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT_BODY
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = C_GRAY_DARK  # gris oscuro, no lila
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r_bold = p.add_run(bold_prefix)
    r_bold.font.name = FONT_BODY
    r_bold.font.bold = True
    r_bold.font.color.rgb = C_BODY
    r_body = p.add_run(text)
    r_body.font.name = FONT_BODY
    r_body.font.color.rgb = C_BODY
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = FONT_BODY
    r.font.color.rgb = C_BODY
    return p

def add_image_centered(doc, img_path, caption):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run()
        r.add_picture(img_path, width=Inches(4.5))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        r_cap = p_cap.add_run(f"Figura: {caption}")
        r_cap.font.name = FONT_BODY
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = C_GRAY_MED
    else:
        print(f"Warning: Imagen no encontrada: {img_path}")

def add_zebra_table(doc, headers, rows):
    """Tabla con encabezado rojo (texto blanco) y filas alternadas F2F2F2 / FFFFFF."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_background(hdr_cells[i], "E30613")
        set_cell_margins(hdr_cells[i])
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT_BODY
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = C_TABLE_ALT if ridx % 2 == 0 else C_TABLE_WHITE
        for c in cells:
            set_cell_background(c, fill)
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT_BODY
            r.font.size = Pt(10)
            r.font.color.rgb = C_BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def compile_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert
        print(f"Converting {docx_path} to PDF...")
        convert(docx_path, pdf_path)
        print("Compilation complete!")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
